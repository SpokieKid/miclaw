from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "opc-course-handbook.docx"
LOGO = ROOT / "yihe.png"
AVATAR = ROOT.parent / "opc_bistro_slides" / "affe-avatar.jpg"

TITLE = "OPC AI 十倍开发效率创业课"
BRAND = "一核学院 · SoloCore Academy"

INTRO_PARAGRAPHS = [
    "这门课程主要面向开发者背景的学员，尤其适合公司内部具备技术背景, 并希望转型为类 OPC 团队的学员。",
    "课程将从开发者视角出发，系统讲清什么是 OPC，以及为什么当下全球范围内越来越多的人正在向 OPC 转型。我们会重点讨论，技术人员如何掌握转型 OPC 所需要的关键能力。也会讲解技术人员应如何如何抓住 AI Coding Tools 带来的独特技术红利，将自己的的工程能力转化为更强的个人竞争力与业务价值。",
    "同时，课程也会正面回应技术人员在转型过程中常见的短板，帮助大家补足设计、沟通、商务、营销和运营等非技术能力，建立更完整的能适应新时代的能力结构。",
    "希望通过这门课，让开发者不仅能把代码写得更快、更好，也能真正看懂趋势、抓住机会，并在新的环境里完成从技术执行者到能独立发现问题, 解决问题的 OPC 型全能人才的升级。",
]

LECTURER_NAME = "张雨霏 · affe"
LECTURER_ROLE = "Einko 联合创始人 / 一核学院 AI Coding 实战课主讲人"
LECTURER_PARAGRAPHS = [
    "张雨霏，affe，全国首家 OPC 社区鸿鹄汇早期入驻团队 Einko 联合创始人，前 Looki Agent 开发工程师，开源社区 Apache RocketMQ Committer，阿里巴巴国际人才孵化中心签约讲师，曾先后在阿里云、腾讯音乐、三星、Splunk、StreamNative 工作。",
    "作为 Claude Code 的早期使用者，拥有长期在 OPC 环境里使用各类 AI 工具为真实业务场景搭建自动化工作流的经验，也踩过大量 AI Coding 的坑。对他来说，真正开心的事情之一，就是把这些经过实战验证的干货分享给更多人。",
]

SESSIONS = [
    {
        "time": "09:00 - 09:45",
        "title": "OPC 政策与创业趋势 + 全国首家 OPC 社区鸿鹄汇创业政策详解",
        "items": [
            "OPC 政策与创业趋势",
            "全国首家 OPC 社区鸿鹄汇的简单介绍",
            "OPC 产生的原因与时代趋势下的业务逻辑改变",
            "技术给开发者带来的红利与挑战",
            "小团队的灵活性与挑战",
            "如何充分利用平台给创作者的激励",
        ],
    },
    {
        "time": "09:50 - 10:35",
        "title": "OPC 的 AI 信息差与利用创业红利：新的范式和思维上的转变",
        "items": [
            "开发者的挑战：做内容以及吸引注意力",
            "个人 IP 之所以如此重要的底层逻辑",
            "用信息差赚钱不要有耻感",
            "如何克服营销羞耻",
            "现金流的重要性",
            "成为 OPC 意味着什么，为什么做好一家公司仍然很困难",
        ],
    },
    {
        "time": "10:45 - 11:30",
        "title": "OPC 的商业行动画布与常见商业模式分享",
        "items": [
            "OPC 的商业模式分析框架",
            "独立开发者",
            "自媒体",
            "外包",
            "知识付费",
            "现金流互联网产品",
            "其他：电商",
        ],
    },
    {
        "time": "13:30 - 14:15",
        "title": "OPC 代码开发：Claude Code、Codex 的基本配置和使用逻辑 (需要用到 Codex 桌面版)",
        "items": [
            "放弃手写代码，认识到代码是 OPC 时代里不稀缺的资源",
            "从 Vibe Coding 氛围编程到 Agentic Engineering 智能体工程化",
            "Agentic Engineering 是什么",
            "什么是 Spec Harness",
            "gstacks 的介绍",
            "Codex 与 Claude Code 的使用场景",
            "Claude Code 最佳使用实践：语音输入、remote-control、loop",
            "常用 Skills 和 agent 设置",
            "如何在手机上随时随地写代码",
        ],
    },
    {
        "time": "14:20 - 15:05",
        "title": "OPC 代码开发：设计到交付全流程 (需要用到 Codex 桌面版)",
        "items": [
            "step1 / htmltodesign：好的设计从模仿开始",
            "Framer / Paper / Variant：有品味的设计",
            "如何判断什么品类值得做，从模仿成熟赛道开始",
            "如何通过 gstacks 启动项目",
            "如何培养 OPC 的设计品味",
        ],
    },
    {
        "time": "15:10 - 15:55",
        "title": "OPC 代码开发：在本地搭建自动化测试框架，以及 Agentic Engineering 处理复杂工程 (需要用到 Codex 桌面版)",
        "items": [
            "为什么需要自动化测试，为什么它对 Coding Agent 如此重要",
            "web-access skills 的安装",
            "自动化测试框架与 Claude Code 的配合使用，包括 Playwright、CDP、chrome-dev-tools",
            "将 Claude Code 融入 CI/CD 工作流中，包括自动化部署以及可观测性",
        ],
    },
    {
        "time": "16:00 - 16:45",
        "title": "OPC 代码开发：多项目 Agent 管理和项目管理",
        "items": [
            "上下文管理的核心要素与工具：使用 Linear 追踪、管理多个项目",
            "使用 Tmux 建立持续的 Coding Session 来完成特定任务",
            "多窗口 Agent Engineering 的开发最佳实践",
            "如何在工程中形成复利，以及如何沉淀 skills",
        ],
    },
    {
        "time": "16:50 - 17:35",
        "title": "OPC 通用智能体：个人智能体环境与知识系统搭建以及其他技能",
        "items": [
            "社交媒体素材库管理和自动发布",
            "Remotion + Screenstudio 自动剪视频工作流",
            "如何建立高质量的信息源和设计风格参考",
            "如何用智能体干杂活：PPT、Excel、Word、Markdown、HTML 操作",
            "初识 SEO、增长和运营",
        ],
    },
]

PREP_ITEMS = [
    "带上一台笔记本电脑，最好是 macOS 操作系统。",
    "闲鱼上准备一个 Codex 的账号，一个月即可，根据不同账号类型花费大概在 12 ~ 140 元左右。",
    "如有条件，可以在闲鱼上准备一个 Claude Code 的账号，一个月大概 150 元左右。",
    "如有条件，在闲鱼上准备一个 Gmail 账号，用来薅各类 AI 工具的羊毛。",
    "Codex 桌面版下载链接：https://developers.openai.com/codex/app",
]


def set_run_font(run, *, name: str = "PingFang SC", size: float = 11, bold: bool = False, color: str | None = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), name)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "E7DDD2", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_heading(doc: Document, text: str, size: float = 16, color: str = "1F1A17"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return p


def add_body(doc_or_cell, text: str, *, size: float = 11, color: str = "4D4035", after: float = 6):
    p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color)
    return p


def add_numbered_item(doc: Document, index: int, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{index}. {text}")
    set_run_font(run, size=10.8, color="4D4035")


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    brand = doc.add_paragraph()
    brand.paragraph_format.space_after = Pt(4)
    if LOGO.exists():
        logo_run = brand.add_run()
        logo_run.add_picture(str(LOGO), width=Cm(3.4))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(BRAND)
    set_run_font(run, size=10.5, bold=True, color="B45309")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(TITLE)
    set_run_font(run, size=22, bold=True, color="1F1A17")

    add_heading(doc, "课程介绍")
    for paragraph in INTRO_PARAGRAPHS:
        add_body(doc, paragraph, size=11.2, after=5)

    add_heading(doc, "实战课主讲人")
    lecturer = doc.add_table(rows=1, cols=2)
    lecturer.autofit = False
    lecturer.columns[0].width = Cm(2.7)
    lecturer.columns[1].width = Cm(13.8)
    avatar_cell, info_cell = lecturer.rows[0].cells
    for cell in (avatar_cell, info_cell):
        set_cell_shading(cell, "FFFDF7")
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    if AVATAR.exists():
        p = avatar_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(AVATAR), width=Cm(1.8))

    p = info_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(LECTURER_NAME)
    set_run_font(run, size=14, bold=True, color="1F1A17")

    p = info_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(LECTURER_ROLE)
    set_run_font(run, size=10.6, bold=True, color="B45309")

    for paragraph in LECTURER_PARAGRAPHS:
        p = info_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(paragraph)
        set_run_font(run, size=10.4, color="4D4035")

    add_heading(doc, "课程时间线")
    timeline = doc.add_table(rows=1, cols=3)
    timeline.autofit = False
    timeline.columns[0].width = Cm(2.7)
    timeline.columns[1].width = Cm(5.2)
    timeline.columns[2].width = Cm(8.1)

    headers = timeline.rows[0].cells
    for cell, text in zip(headers, ("时间", "主题", "内容要点")):
        set_cell_shading(cell, "F7E9D7")
        set_cell_border(cell, color="E1C59D")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=10.5, bold=True, color="7C3E09")

    for session in SESSIONS:
        row = timeline.add_row().cells
        for cell in row:
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        p = row[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(session["time"])
        set_run_font(run, size=10.3, bold=True, color="1F1A17")

        p = row[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(session["title"])
        set_run_font(run, size=10.5, bold=True, color="1F1A17")

        detail_p = row[2].paragraphs[0]
        detail_p.paragraph_format.space_after = Pt(0)
        first = True
        for item in session["items"]:
            p = detail_p if first else row[2].add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.first_line_indent = Cm(-0.28)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(f"• {item}")
            set_run_font(run, size=9.7, color="4D4035")
            first = False

    add_heading(doc, "课前准备")
    for idx, item in enumerate(PREP_ITEMS, start=1):
        add_numbered_item(doc, idx, item)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
